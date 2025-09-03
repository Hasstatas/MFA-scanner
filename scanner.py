import os
from pathlib import Path
import pytesseract
from pytesseract import TesseractNotFoundError
from PIL import Image
from strategies import load_strategies
from reports.report_service import generate_pdf
from typing import Optional
try:
    import fitz  
except Exception:
    fitz = None
try:
    from docx import Document as DocxDocument
except Exception:
    DocxDocument = None

# ------------- Tesseract (bundled or system) -------------
DEFAULT_TESS_PATHS = [
    # Bundled portable exe (preferred)
    str(Path(__file__).parent / "vendor" / "tesseract" / "tesseract.exe"),
    str(Path.cwd() / "vendor" / "tesseract" / "tesseract.exe"),
    # Common Windows installs
    r"C:\Program Files\Tesseract-OCR\tesseract.exe",
    r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
]

def configure_tesseract() -> bool:

    env_cmd = os.environ.get("TESSERACT_CMD")
    candidates = [env_cmd] if env_cmd else []
    candidates.extend(DEFAULT_TESS_PATHS)

    for c in candidates:
        if not c:
            continue
        p = Path(c)
        if p.exists():
            pytesseract.pytesseract.tesseract_cmd = str(p)
            tessdata_dir = p.parent / "tessdata"
            os.environ["TESSDATA_PREFIX"] = str(tessdata_dir if tessdata_dir.exists() else p.parent)
            print(f"[OCR] Using tesseract: {p}")
            print(f"[OCR] TESSDATA_PREFIX: {os.environ['TESSDATA_PREFIX']}")
            return True

    print("[OCR] Tesseract not found. Set TESSERACT_CMD or add vendor/tesseract/tesseract.exe")
    return False

# ----------------------- UI helpers -----------------------
def choose_one(title: str, options: list[str]) -> str:
    """Prompt user to choose exactly one option by number."""
    print(title)
    for i, o in enumerate(options, 1):
        print(f"{i}. {o}")
    while True:
        raw = input("Select ONE option (e.g. 1): ").strip()
        if raw.isdigit():
            i = int(raw)
            if 1 <= i <= len(options):
                return options[i - 1]
        print("Invalid selection. Please enter a single number from the list.")

def list_evidence(folder: Path):
    # Supported evidence types
    image_exts = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"}
    other_exts = {".pdf", ".txt", ".docx"}
    exts = image_exts | other_exts
    return sorted([p for p in folder.iterdir() if p.is_file() and p.suffix.lower() in exts])

def choose_one_evidence(evidence_dir: Path) -> Path | None:
    """
    Show evidence files in evidence/ and prompt for exactly one.
    Also accepts a full path to a file outside the folder.
    """
    files = list_evidence(evidence_dir)
    if files:
        print("\nEvidence files (from ./evidence):")
        for i, p in enumerate(files, 1):
            print(f"{i}. {p.name}")
    else:
        print("No supported files found in ./evidence (you can paste a full file path).")

    while True:
        resp = input("Pick ONE evidence file by number, or paste a full path: ").strip()
        if not resp:
            print("Please choose a number or paste a file path.")
            continue
        if resp.isdigit() and files:
            i = int(resp)
            if 1 <= i <= len(files):
                return files[i - 1]
            print("Number out of range.")
            continue
        pr = Path(resp)
        if pr.exists() and pr.is_file():
            return pr

        print("Invalid selection. Try again.")

# ------------------- text + preview extraction -------------------
def _ocr_image(path: Path) -> str:
    try:
        img = Image.open(path)
        return pytesseract.image_to_string(img)
    except TesseractNotFoundError:
        print("⚠️  Tesseract not found – continuing without OCR. 'Extract' will be blank.")
        return ""
    except Exception as e:
        print(f"⚠️  OCR failed on {path.name}: {e} — continuing with blank Extract.")
        return ""

def _extract_pdf(path: Path, previews_dir: Path) -> tuple[str, Optional[Path]]:
    if not fitz:
        print("⚠️  PyMuPDF (pymupdf) not installed; cannot parse PDF. `pip install pymupdf`")
        return "", None
    try:
        doc = fitz.open(str(path))
        text = ""
        for page in doc:
            text += page.get_text() or ""
        preview_path = None
        if len(doc) > 0:
            page = doc[0]
            pix = page.get_pixmap()
            previews_dir.mkdir(parents=True, exist_ok=True)
            preview_path = previews_dir / (path.stem + "_page1.png")
            pix.save(str(preview_path))
        doc.close()
        return text, preview_path
    except Exception as e:
        print(f"⚠️  PDF parse failed on {path.name}: {e}")
        return "", None

def _extract_docx(path: Path) -> str:
    if not DocxDocument:
        print("⚠️  python-docx not available for DOCX extract (should already be installed).")
        return ""
    try:
        doc = DocxDocument(str(path))
        parts = []
        for p in doc.paragraphs:
            if p.text:
                parts.append(p.text)
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    if cell.text:
                        parts.append(cell.text)
        return "\n".join(parts)
    except Exception as e:
        print(f"⚠️  DOCX extract failed on {path.name}: {e}")
        return ""

def _extract_txt(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8", errors="ignore")
    except Exception as e:
        print(f"⚠️  TXT read failed on {path.name}: {e}")
        return ""

def extract_text_and_preview(path: Path, previews_dir: Path) -> tuple[str, Optional[Path]]:
    ext = path.suffix.lower()
    if ext in {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"}:
        return _ocr_image(path), path  # preview is original image
    if ext == ".pdf":
        return _extract_pdf(path, previews_dir)
    if ext == ".docx":
        return _extract_docx(path), None
    if ext == ".txt":
        return _extract_txt(path), None
    print(f"⚠️  Unsupported evidence type: {ext}")
    return "", None


# ------------------------ Main --------------------------
def main():
    configure_tesseract() 

    user_id = input("Enter your username: ").strip()

    strategies = load_strategies()
    if not strategies:
        print("No strategies found.")
        return

    def desc(s):
        try:
            return s.description()
        except Exception:
            return ""

    names_for_menu = [f"{s.name} — {desc(s)}" for s in strategies]
    chosen_label = choose_one("\nAvailable strategies:", names_for_menu)
    chosen_name = chosen_label.split(" — ")[0]
    chosen_strategy = next((s for s in strategies if s.name == chosen_name), None)
    if not chosen_strategy:
        print("Not a valid selection. Exiting.")
        return

    print(f"\n📋 Strategy selected: {chosen_strategy.name}\n")

    evidence_dir = Path("evidence")
    evidence_dir.mkdir(parents=True, exist_ok=True)

    evidence_path = choose_one_evidence(evidence_dir)
    if not evidence_path:
        print("No evidence selected.")
        return

    out_dir = Path("reports_out")
    out_dir.mkdir(parents=True, exist_ok=True)
    previews_dir = Path("previews")  
    template_path = "templates/report_template.docx"

    print(f"\n📄  Evidence: {evidence_path.name}")
    text, preview = extract_text_and_preview(evidence_path, previews_dir)
    preview_text = (text[:300] + "…") if len(text) > 300 else text
    print("📝 Extracted Text:", preview_text)

    generated = []

    try:
        if hasattr(chosen_strategy, "emit_hits"):
            rows = chosen_strategy.emit_hits(text) or []  
            if not rows:
                print("No findings for this strategy.")
            for idx, r in enumerate(rows, start=1):
                uid = _safe_uid(user_id)
                data = {
                    "UniqueID": uid,
                    "UserID": user_id,
                    "Evidence": str(evidence_path),
                    "Evidence Preview": str(preview) if preview else "",
                    "Strategy": chosen_strategy.name,
                    "TestID": r.get("test_id", ""),
                    "Sub-Strategy": r.get("sub_strategy", ""),
                    "ML Level": r.get("detected_level", ""),
                    "Pass/Fail": r.get("pass_fail", ""),
                    "Priority": r.get("priority", ""),
                    "Recommendation": r.get("recommendation", ""),
                    "Evidence Extract": "; ".join(r.get("evidence", [])),
                    "Description": r.get("description", ""),
                    "Confidence": r.get("confidence", ""),
                }
                pdf = generate_pdf(
                    data,
                    template_path=template_path,
                    output_dir=str(out_dir),
                    base_dir="."
                )
                print(f"   ✅ {chosen_strategy.name} → {pdf}")
                generated.append(pdf)
        else:
            hits = chosen_strategy.match(text) or []
            if not hits:
                print("No findings for this strategy.")
            else:
                uid = _safe_uid(f"{user_id}-{chosen_strategy.name}-{evidence_path.stem}-HITS")
                data = {
                    "UniqueID": uid,
                    "UserID": user_id,
                    "Evidence": str(evidence_path),
                    "Evidence Preview": str(preview) if preview else "",
                    "Strategy": chosen_strategy.name,
                    "TestID": "",
                    "Sub-Strategy": "",
                    "ML Level": "",
                    "Pass/Fail": "",
                    "Priority": "",
                    "Recommendation": "",
                    "Evidence Extract": ", ".join(hits),
                }
                pdf = generate_pdf(
                    data,
                    template_path=template_path,
                    output_dir=str(out_dir),
                    base_dir="."
                )
                print(f"   ✅ {chosen_strategy.name} → {pdf}")
                generated.append(pdf)
    except Exception as e:
        print(f"   ❌ {chosen_strategy.name} failed: {e}")

    if generated:
        print("\n📄 Generated reports:")
        for p in generated:
            print(" -", p)
    else:
        print("\nNo findings → no reports generated.")


# -------------------- utilities --------------------
def _safe_uid(s: str) -> str:
    keep = []
    for ch in s:
        if ch.isalnum() or ch in ("-", "_"):
            keep.append(ch)
        else:
            keep.append("-")
    return "".join(keep).strip("-") or "report"


if __name__ == "__main__":
=======
# scanner.py
import os
import csv
from pathlib import Path
from typing import List
import pytesseract
from pytesseract import TesseractNotFoundError
from PIL import Image, UnidentifiedImageError

# If PATH sometimes fails, uncomment and set your path explicitly:
# pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# Optional but recommended for better OCR
try:
    import cv2
    _HAS_CV2 = True
except Exception:
    _HAS_CV2 = False

from strategies import load_strategies  # your plugin loader

# ---------- Folder mapping for the 8 Essential Eight strategies ----------
STRAT_DIR_MAP = {
    "application control": "application_control",
    "restrict admin privileges": "restrict_admin_privileges",
    "patch applications": "patch_applications",
    "patch operating systems": "patch_operating_systems",
    "configure microsoft office macro settings": "configure_macro_settings",
    "multi-factor authentication": "multi_factor_authentication",
    "regular backups": "regular_backups",
    "user application hardening": "user_application_hardening",
}

# ---------- File type support ----------
SUPPORTED_IMAGE_EXTS = (".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp")
SUPPORTED_TEXT_EXTS  = (".txt", ".log", ".reg", ".csv", ".ini", ".json", ".xml", ".htm", ".html")

# ---------- UI helpers ----------
def choose_from_menu(title: str, options: List[str]) -> List[str]:
    print(title)
    for i, o in enumerate(options, 1):
        print(f"{i}. {o}")
    raw = input("Select the strategy by number(s) e.g. 1,3,5 for multiple strategies: ")
    idxs = []
    for tok in raw.split(","):
        tok = tok.strip()
        if tok.isdigit() and 1 <= int(tok) <= len(options):
            idxs.append(int(tok) - 1)
    return [options[i] for i in idxs]

# ---------- OCR helpers ----------
def _ocr_with_pillow(path: Path) -> str:
    """Plain PIL -> pytesseract (fallback)."""
    try:
        with Image.open(path) as img:
            return pytesseract.image_to_string(img, config="--psm 6")
    except (UnidentifiedImageError, OSError, TesseractNotFoundError):
        return ""

def _ocr_with_cv2(path: Path) -> str:
    """OpenCV preprocessing for clearer OCR."""
    if not _HAS_CV2:
        return _ocr_with_pillow(path)

    try:
        img = cv2.imread(str(path), cv2.IMREAD_COLOR)
        if img is None:
            return ""

        # Grayscale
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)

        # Upscale a bit to help OCR on UI screenshots
        gray = cv2.resize(gray, None, fx=1.5, fy=1.5, interpolation=cv2.INTER_CUBIC)

        # Boost contrast/brightness
        gray = cv2.convertScaleAbs(gray, alpha=2.0, beta=15)

        # Binarize adaptively to handle uneven backgrounds
        thr = cv2.adaptiveThreshold(
            gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
            cv2.THRESH_BINARY, 35, 11
        )

        return pytesseract.image_to_string(thr, config="--psm 6")
    except TesseractNotFoundError:
        return ""
    except Exception:
        # Last-ditch fallback
        return _ocr_with_pillow(path)

def run_ocr(path: Path) -> str:
    """Try OpenCV pre-processing first, then fallback to plain OCR."""
    text = _ocr_with_cv2(path)
    if text and text.strip():
        return text
    return _ocr_with_pillow(path)

# ---------- Content extraction ----------
def read_text_file(path: Path) -> str:
    """Read small text-like files into a single string."""
    try:
        return path.read_text(encoding="utf-8", errors="ignore")
    except Exception:
        try:
            return path.read_text(errors="ignore")
        except Exception:
            return ""

def extract_text(path: Path) -> str:
    """Choose OCR for images, direct read for text files."""
    ext = path.suffix.lower()
    if ext in SUPPORTED_IMAGE_EXTS:
        return run_ocr(path)
    if ext in SUPPORTED_TEXT_EXTS:
        return read_text_file(path)
    return ""  # unsupported

def list_supported_files(folder: Path) -> List[Path]:
    """Recursively gather supported files under folder."""
    if not folder.exists():
        return []
    out = []
    for p in folder.rglob("*"):  # RECURSIVE now
        if not p.is_file():
            continue
        ext = p.suffix.lower()
        if ext in SUPPORTED_IMAGE_EXTS or ext in SUPPORTED_TEXT_EXTS:
            out.append(p)
    return sorted(out, key=lambda x: x.name.lower())

# ---------- Main ----------
def main():
    # 1) user id
    user_id = input("Enter your username: ").strip()

    # 2) load strategies
    strategies = load_strategies()
    if not strategies:
        print("No strategies found")
        return

    # 3) pick strategies
    names_for_menu = [f"{s.name} — {s.description()}" for s in strategies]
    chosen_menu = choose_from_menu("\nAvailable strategies:", names_for_menu)

    chosen_names = {n.split(" — ")[0] for n in chosen_menu}
    chosen_strategies = [s for s in strategies if s.name in chosen_names]

    if not chosen_strategies:
        print("Not a valid selection. Exiting the application.")
        return

    print("\n📋 Scanning using strategies:", ", ".join(s.name for s in chosen_strategies), "\n")

    # 4) where screenshots/files live
    # Put files in: screenshots/<mapped_subdir>/...
    base_dir = Path("screenshots")

    # 5) CSV header
    report_rows = [(
        "UserID", "Image", "Strategy", "TestID", "Sub-Strategy",
        "ML Level", "Pass/Fail", "Priority", "Recommendation", "Evidence Extract"
    )]

    # 6) scan per strategy
    for strat in chosen_strategies:
        strat_key = strat.name.lower().strip()
        strat_sub = STRAT_DIR_MAP.get(strat_key, None)
        preferred = (base_dir / strat_sub) if strat_sub else base_dir
        fallback  = base_dir

        files = list_supported_files(preferred)
        using_dir = preferred
        if not files:
            files = list_supported_files(fallback)
            using_dir = fallback

        if not files:
            print(f"⚠️  No files found for '{strat.name}' in '{preferred}' or '{fallback}'. Skipping.")
            continue

        print(f"\n🔎 Strategy: {strat.name}")
        print(f"   Using inputs from: {using_dir}")

        for fpath in files:
            print(f"📄 {fpath.name}:")
            raw_text = extract_text(fpath)

            if not raw_text.strip():
                print("   (no readable text found)\n")
                # record a row so you see the file in the CSV
                report_rows.append((
                    user_id, fpath.name, strat.name, "", "", "",
                    "NO_TEXT", "Low",
                    "OCR could not read this file. Try a clearer screenshot.",
                    ""
                ))
                continue

            preview = (raw_text[:200] + "…") if len(raw_text) > 200 else raw_text
            print("📝 Extracted:", preview.replace("\n", " ")[:200], "\n")

            rows_added = 0  # track whether any hits were written

            if hasattr(strat, "emit_hits"):
                rows = strat.emit_hits(raw_text, source_file=fpath.name)
                for r in rows:
                    report_rows.append((
                        user_id,
                        fpath.name,
                        strat.name,
                        r.get("test_id", ""),
                        r.get("sub_strategy", ""),
                        r.get("detected_level", ""),
                        r.get("pass_fail", ""),
                        r.get("priority", ""),
                        r.get("recommendation", ""),
                        "; ".join(r.get("evidence", [])),
                    ))
                    rows_added += 1
            else:
                hits = strat.match(raw_text)
                if hits:
                    report_rows.append((
                        user_id, fpath.name, strat.name, "", "", "",
                        "HIT", "Medium",
                        "Heuristic match.",
                        ", ".join(hits)
                    ))
                    rows_added += 1

            # If no matches, write a NO_MATCH row so the file appears in the report
            if rows_added == 0:
                report_rows.append((
                    user_id, fpath.name, strat.name, "", "", "",
                    "NO_MATCH", "Low",
                    "No rule matched this file.",
                    ""
                ))

    # 7) save
    try:
        with open("scan_report.csv", "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerows(report_rows)
        print("\n✅ Report saved as: scan_report.csv")
    except PermissionError:
        temp_name = "scan_report_temp.csv"
        with open(temp_name, "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerows(report_rows)
        print(f"\n⚠️  'scan_report.csv' was locked. Saved as: {temp_name}")

if __name__ == "__main__":
    # Helpful check: print tesseract path/version once
    try:
        ver = pytesseract.get_tesseract_version()
        print(f"(i) Tesseract found: {ver}")
    except Exception:
        print("(i) Tesseract not detected by pytesseract. Ensure the Tesseract OCR engine is installed and on PATH.")
    main()
