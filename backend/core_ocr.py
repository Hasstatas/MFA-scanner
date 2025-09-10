from pathlib import Path
from typing import Optional, Tuple
from shutil import which
import os
import pytesseract
from pytesseract import TesseractNotFoundError
from PIL import Image, UnidentifiedImageError

try:
    import cv2
    _HAS_CV2 = True
except Exception:
    _HAS_CV2 = False

try:
    import fitz  
except Exception:
    fitz = None

try:
    from docx import Document as DocxDocument
except Exception:
    DocxDocument = None

# Filetype support
SUPPORTED_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"}
SUPPORTED_TEXT_EXTS  = {".txt", ".log", ".reg", ".csv", ".ini", ".json", ".xml", ".htm", ".html"}
SUPPORTED_DOC_EXTS   = {".docx", ".pdf"}
SUPPORTED_ALL_EXTS   = SUPPORTED_IMAGE_EXTS | SUPPORTED_TEXT_EXTS | SUPPORTED_DOC_EXTS

def configure_tesseract() -> None:
    cmd = os.environ.get("TESSERACT_CMD") or which("tesseract")
    if not cmd:
        print("[OCR] Tesseract not found. Install it or set TESSERACT_CMD.")
        return
    pytesseract.pytesseract.tesseract_cmd = cmd
    print(f"[OCR] Using tesseract: {cmd}")
    if "TESSDATA_PREFIX" in os.environ:
        print(f"[OCR] TESSDATA_PREFIX: {os.environ['TESSDATA_PREFIX']}")

def _ocr_with_pillow(path: Path) -> str:
    try:
        with Image.open(path) as img:
            return pytesseract.image_to_string(img, config="--psm 6")
    except (UnidentifiedImageError, OSError, TesseractNotFoundError):
        return ""

def _ocr_with_cv2(path: Path) -> str:
    if not _HAS_CV2:
        return _ocr_with_pillow(path)
    try:
        import cv2
        img = cv2.imread(str(path), cv2.IMREAD_COLOR)
        if img is None:
            return ""
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        gray = cv2.resize(gray, None, fx=1.5, fy=1.5, interpolation=cv2.INTER_CUBIC)
        gray = cv2.convertScaleAbs(gray, alpha=2.0, beta=15)
        thr = cv2.adaptiveThreshold(
            gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
            cv2.THRESH_BINARY, 35, 11
        )
        return pytesseract.image_to_string(thr, config="--psm 6")
    except TesseractNotFoundError:
        return ""
    except Exception:
        return _ocr_with_pillow(path)

def ocr_image(path: Path) -> str:
    txt = _ocr_with_cv2(path)
    return txt if txt.strip() else _ocr_with_pillow(path)

def extract_text_and_preview(path: Path, previews_dir: Path) -> Tuple[str, Optional[Path]]:
    """
    Returns (text, preview_path_or_None)
    - For images: run OCR and ALWAYS create a PNG preview
    - For PDFs: text via PyMuPDF + page 1 preview
    - For DOCX: extract text via python-docx
    - For Text-like files: return file text 
    """
    ext = path.suffix.lower()

    # For Images
    if ext in SUPPORTED_IMAGE_EXTS:
        text = ocr_image(path)

        # Always have a PNG preview
        previews_dir.mkdir(parents=True, exist_ok=True)
        preview_png = previews_dir / f"{path.stem}.png"

        try:
            with Image.open(path) as img:
                # Convert to RGB to avoid palette/transparency issues
                if img.mode in ("P", "RGBA"):
                    img = img.convert("RGB")
                img.save(preview_png, format="PNG")
            return text, preview_png
        except Exception:
            # If conversion fails, return text but no preview
            return text, None

    # For PDFs 
    if ext == ".pdf":
        if not fitz:
            print("⚠ PyMuPDF not installed; cannot parse PDF. pip install pymupdf")
            return "", None
        try:
            doc = fitz.open(str(path))
            text = "".join(page.get_text() or "" for page in doc)
            preview_path = None
            if len(doc) > 0:
                pix = doc[0].get_pixmap()
                previews_dir.mkdir(parents=True, exist_ok=True)
                preview_path = previews_dir / f"{path.stem}_page1.png"
                pix.save(str(preview_path))
            doc.close()
            return text, preview_path
        except Exception as e:
            print(f"⚠ PDF parse failed on {path.name}: {e}")
            return "", None

    # For DOCX
    if ext == ".docx":
        if not DocxDocument:
            print("⚠ python-docx not available.")
            return "", None
        try:
            doc = DocxDocument(str(path))
            parts = []
            for p in doc.paragraphs:
                if p.text:
                    parts.append(p.text)
            for tbl in doc.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        if cell.text:
                            parts.append(cell.text)
            return "\n".join(parts), None
        except Exception as e:
            print(f"⚠ DOCX parse failed on {path.name}: {e}")
            return "", None

    # For Text-like files
    if ext in SUPPORTED_TEXT_EXTS:
        try:
            return path.read_text(encoding="utf-8", errors="ignore"), None
        except Exception:
            try:
                return path.read_text(errors="ignore"), None
            except Exception:
                return "", None
                
    # If it does not meet any of the file type
    return "", None