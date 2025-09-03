from __future__ import annotations

import os
import uuid
import subprocess
from datetime import datetime
from pathlib import Path
from typing import Mapping, Any, Optional, Dict, Tuple

from docx import Document
from docx.shared import Inches


def generate_pdf(
    data: Mapping[str, Any],
    *,
    template_path: os.PathLike | str = "templates/report_template.docx",
    output_dir: os.PathLike | str = "reports_out",
    base_dir: os.PathLike | str = ".",
    image_marker: str = "[Embed evidence here]",   
    unique_id_override: Optional[str] = None,
) -> Path:
    """
    Render a single PDF from the in-memory mapping produced by the OCR/rules step.

    Expected keys in `data` (case/spacing tolerant):
      UniqueID or UserID -> becomes UniqueID in template
      Evidence -> path to original evidence file 
      Evidence Preview (optional) -> path to an image to embed 
      Strategy, TestID, Sub-Strategy, ML Level, Pass/Fail, Priority,
      Recommendation -> Recommendations, Evidence Extract -> Extract
      Description 
      Confidence 

    Returns: Path to the generated PDF.
    """
    mapping, embed_path, unique_id = _map_to_placeholders(data, Path(base_dir))

    if unique_id_override:
        unique_id = unique_id_override
        mapping["UniqueID"] = unique_id
        mapping["Unique ID"] = unique_id

    tpath = Path(template_path)
    if not tpath.exists():
        raise FileNotFoundError(f"Template not found: {tpath}")

    outdir = Path(output_dir)
    outdir.mkdir(parents=True, exist_ok=True)
    pdf_path = outdir / f"{unique_id}.pdf"
    doc = Document(str(tpath))
    _replace_braced_placeholders_everywhere(doc, mapping)
    _replace_xml_text_everywhere(doc, mapping)
    if embed_path:
        if not _insert_image_at_marker(doc, image_marker, embed_path, width_inches=6.0):
            _insert_image_at_marker(doc, "[Embed screenshot here]", embed_path, width_inches=6.0)
    else:
        _remove_markers_everywhere(doc, ["[Embed evidence here]", "[Embed screenshot here]"])
    filled = pdf_path.with_suffix(".filled.docx")
    doc.save(str(filled))
    _convert_docx_to_pdf(filled, pdf_path)
    try:
        filled.unlink()
    except Exception:
        pass

    return pdf_path


# ---------- Mapping (OCR dict -> template placeholders) ----------

def _normalize_keys(d: Mapping[str, Any]) -> Dict[str, str]:
    norm: Dict[str, str] = {}
    for k, v in d.items():
        key = " ".join(str(k).strip().lower().replace("_", " ").replace("-", " ").replace("/", " ").split())
        norm[key] = "" if v is None else str(v)
    return norm

def _pick(norm: Dict[str, str], *names: str) -> str:
    for n in names:
        key = " ".join(n.strip().lower().split())
        if key in norm:
            return norm[key]
    return ""

def _map_to_placeholders(data: Mapping[str, Any], base_dir: Path) -> Tuple[Dict[str, str], Optional[Path], str]:
    n = _normalize_keys(data)

    # Inputs (tolerant keys)
    unique_id = _pick(n, "uniqueid", "unique id", "userid", "user id") or str(uuid.uuid4())
    strategy  = _pick(n, "strategy")
    testid    = _pick(n, "testid", "test id")
    substrat  = _pick(n, "sub-strategy", "sub strategy")
    level     = _pick(n, "ml level", "level")
    passfail  = _pick(n, "pass/fail", "pass fail")
    priority  = _pick(n, "priority")
    rec       = _pick(n, "recommendation", "recommendations")
    extract   = _pick(n, "evidence extract", "extract")
    descr     = _pick(n, "description")
    confidence = _pick(n, "confidence")  

    # Evidence paths
    evidence_path_str = _pick(n, "evidence", "evidence path", "file", "file path", "filepath", "image", "screenshot")
    preview_path_str  = _pick(n, "evidence preview", "preview", "embed path")

    # Resolve paths
    embed_path: Optional[Path] = None  
    file_name = ""
    if evidence_path_str:
        ep = Path(evidence_path_str)
        if not ep.is_absolute():
            ep = base_dir / ep
        file_name = ep.name
    if preview_path_str:
        pp = Path(preview_path_str)
        if not pp.is_absolute():
            pp = base_dir / pp
        if pp.exists():
            embed_path = pp
    else:
        if evidence_path_str:
            ep = Path(evidence_path_str)
            if not ep.is_absolute():
                ep = base_dir / ep
            if ep.exists() and ep.suffix.lower() in {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"}:
                embed_path = ep

    # Mapping to template placeholders (support a couple of variants)
    mapping: Dict[str, str] = {
        "UniqueID": unique_id,
        "Unique ID": unique_id,
        "UserID": unique_id,            

        "Strategy": strategy,
        "Test_id": testid,
        "Sub-Strategy": substrat,

        "level": level,
        "Level": level,

        "Pass/Fail": passfail,
        "Priority": priority,

        "Recommendations": rec,

        "extract": extract,
        "Extract": extract,

        "Description": descr,
        "description": descr,
        "Confidence": confidence or "",

        "file name": file_name,
        "File Name": file_name,

        "Date Generated": datetime.now().strftime("%d %b %Y"),
    }
    _expand_placeholder_variants(mapping)
    return mapping, embed_path, unique_id


# ---------- DOCX helpers ----------

def _iter_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p

def _replace_in_runs(paragraph, mapping: Mapping[str, Any]) -> bool:
    changed = False
    for run in paragraph.runs:
        txt = run.text
        new = txt
        for k, v in mapping.items():
            new = new.replace("{" + k + "}", str(v))
        if new != txt:
            run.text = new
            changed = True
    return changed

def _rebuild_paragraph_text(paragraph, mapping: Mapping[str, Any]) -> None:
    full = "".join(run.text for run in paragraph.runs)
    repl = full
    for k, v in mapping.items():
        repl = repl.replace("{" + k + "}", str(v))
    if repl != full:
        for r in paragraph.runs:
            r.text = ""
        paragraph.add_run(repl)

def _replace_braced_placeholders_everywhere(doc, mapping: Mapping[str, Any]) -> None:
    for p in _iter_paragraphs(doc):
        if not _replace_in_runs(p, mapping):
            _rebuild_paragraph_text(p, mapping)

def _replace_xml_text_everywhere(doc, mapping: Mapping[str, Any]) -> None:
    """
    Replace {tokens} in all <w:t> text nodes across main doc part,
    headers, and footers. Works even with older python-docx (no namespaces kwarg).
    """
    def replace_in_part(part):
        root = part.element
        texts = []
        try:
            ns = getattr(root, "nsmap", None)
            if ns:
                texts = root.xpath(".//w:t", namespaces=ns)
        except TypeError:
            texts = []
        if not texts:
            texts = root.xpath(".//*[local-name()='t']")

        for t in texts:
            old = t.text or ""
            new = old
            for k, v in mapping.items():
                new = new.replace("{" + k + "}", str(v))
            if new != old:
                t.text = new

    replace_in_part(doc.part)
    for section in doc.sections:
        try:
            if section.header:
                replace_in_part(section.header.part)
        except Exception:
            pass
        try:
            if section.footer:
                replace_in_part(section.footer.part)
        except Exception:
            pass

def _insert_image_at_marker(doc, marker: str, image_path: os.PathLike | str, width_inches: float = 6.0) -> bool:
    ip = Path(image_path)
    if not ip.exists():
        return False
    target = None
    for p in _iter_paragraphs(doc):
        if marker in "".join(run.text for run in p.runs):
            target = p
            break
    if not target:
        return False
    for r in target.runs:
        r.text = r.text.replace(marker, "")
    run = target.add_run()
    try:
        run.add_picture(str(ip), width=Inches(width_inches))
    except AttributeError:
        doc.add_picture(str(ip), width=Inches(width_inches))
    return True

def _convert_docx_to_pdf(input_docx: Path, output_pdf: Path) -> None:
    """
    Try docx2pdf (uses Word on Windows/macOS). If unavailable, fall back to LibreOffice.
    """
    # Preferred: docx2pdf
    try:
        from docx2pdf import convert
        convert(str(input_docx), str(output_pdf))
        return
    except Exception:
        pass

    # Fallback: LibreOffice
    try:
        out_dir = str(output_pdf.parent.resolve())
        subprocess.run(
            ["soffice", "--headless", "--convert-to", "pdf", "--outdir", out_dir, str(input_docx.resolve())],
            check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE
        )
        expected = output_pdf.with_suffix(".pdf")
        if expected.exists() and expected != output_pdf:
            expected.replace(output_pdf)
    except Exception as e:
        raise RuntimeError(
            "PDF conversion failed. Install Microsoft Word for docx2pdf or LibreOffice for fallback."
        ) from e
        
def _remove_markers_everywhere(doc, markers: list[str]) -> None:
    for p in _iter_paragraphs(doc):
        full = "".join(r.text for r in p.runs)
        new_full = full
        for m in markers:
            new_full = new_full.replace(m, "")
        if new_full != full:
            for r in p.runs:
                r.text = ""
            p.add_run(new_full)

    def scrub_part(part):
        root = part.element
        texts = []
        try:
            ns = getattr(root, "nsmap", None)
            if ns:
                texts = root.xpath(".//w:t", namespaces=ns)
        except TypeError:
            texts = []
        if not texts:
            texts = root.xpath(".//*[local-name()='t']")
        for t in texts:
            old = t.text or ""
            new = old
            for m in markers:
                new = new.replace(m, "")
            if new != old:
                t.text = new

    scrub_part(doc.part)
    for section in doc.sections:
        try:
            if section.header:
                scrub_part(section.header.part)
        except Exception:
            pass
        try:
            if section.footer:
                scrub_part(section.footer.part)
        except Exception:
            pass

# ---------- Tolerant placeholder variants ----------

def _expand_placeholder_variants(mapping: Dict[str, str]) -> None:
    """
    Make our {token} replacement tolerant to:
      - non-ASCII hyphens/dashes (‐, -, –, —)
      - optional spaces inside braces: { Token } as well as {Token}
    This only adds alias keys; it does NOT change original keys/values.
    """
    hyphens = ["-", "\u2010", "\u2011", "\u2013", "\u2014"] 
    to_add: Dict[str, str] = {}

    for k, v in list(mapping.items()):
        spaced = f" {k} "
        if spaced not in mapping:
            to_add[spaced] = v
        if "-" in k:
            for h in hyphens:
                if h == "-":
                    continue
                k_dash = k.replace("-", h)
                if k_dash not in mapping:
                    to_add[k_dash] = v
                spaced_dash = f" {k_dash} "
                if spaced_dash not in mapping:
                    to_add[spaced_dash] = v

    mapping.update(to_add)
