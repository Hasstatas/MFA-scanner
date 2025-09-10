import os
from pathlib import Path
import pytesseract
from pytesseract import TesseractNotFoundError
from PIL import Image
from strategies import load_strategies
from reports.report_service import generate_pdf
from typing import Optional
try:
    import fitz  
except Exception:
    fitz = None
try:
    from docx import Document as DocxDocument
except Exception:
    DocxDocument = None
from shutil import which 

# ------------- Tesseract (not bundled) -------------

def configure_tesseract() -> bool:
    env_cmd = os.environ.get("TESSERACT_CMD")
    cmd = env_cmd or which("tesseract")
    if not cmd:
        print("[OCR] Tesseract not found. Install it and ensure it's on PATH, or set TESSERACT_CMD.")
        return False

    import pytesseract
    pytesseract.pytesseract.tesseract_cmd = cmd

    print(f"[OCR] Using tesseract: {cmd}")
    if "TESSDATA_PREFIX" in os.environ:
        print(f"[OCR] TESSDATA_PREFIX: {os.environ['TESSDATA_PREFIX']}")
    return True

# ----------------------- UI helpers -----------------------
def choose_one(title: str, options: list[str]) -> str:
    """Prompt user to choose exactly one option by number."""
    print(title)
    for i, o in enumerate(options, 1):
        print(f"{i}. {o}")
    while True:
        raw = input("Select ONE option (e.g. 1): ").strip()
        if raw.isdigit():
            i = int(raw)
            if 1 <= i <= len(options):
                return options[i - 1]
        print("Invalid selection. Please enter a single number from the list.")

def list_evidence(folder: Path):
    # Supported evidence types
    image_exts = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"}
    other_exts = {".pdf", ".txt", ".docx"}
    exts = image_exts | other_exts
    # allow the tool to find files within folders as well 
    return sorted(
        [p for p in folder.rglob("*") if p.is_file() and p.suffix.lower() in exts],
        key=lambda p: p.as_posix().lower()
    )

def choose_one_evidence(evidence_dir: Path) -> Path | None:
    """
    Show evidence files in evidence/ and prompt for exactly one.
    Also accepts a full path to a file outside the folder.
    """
    files = list_evidence(evidence_dir)
    if files:
        print("\nEvidence files (from ./evidence):")
        for i, p in enumerate(files, 1):
            print(f"{i}. {p.name}")
    else:
        print("No supported files found in ./evidence (you can paste a full file path).")

    while True:
        resp = input("Pick ONE evidence file by number, or paste a full path: ").strip()
        if not resp:
            print("Please choose a number or paste a file path.")
            continue
        if resp.isdigit() and files:
            i = int(resp)
            if 1 <= i <= len(files):
                return files[i - 1]
            print("Number out of range.")
            continue
        pr = Path(resp)
        if pr.exists() and pr.is_file():
            return pr

        print("Invalid selection. Try again.")

# ------------------- text + preview extraction -------------------
def _ocr_image(path: Path) -> str:
    try:
        img = Image.open(path)
        return pytesseract.image_to_string(img)
    except TesseractNotFoundError:
        print("⚠️  Tesseract not found – continuing without OCR. 'Extract' will be blank.")
        return ""
    except Exception as e:
        print(f"⚠️  OCR failed on {path.name}: {e} — continuing with blank Extract.")
        return ""

def _extract_pdf(path: Path, previews_dir: Path) -> tuple[str, Optional[Path]]:
    if not fitz:
        print("⚠️  PyMuPDF (pymupdf) not installed; cannot parse PDF. `pip install pymupdf`")
        return "", None
    try:
        doc = fitz.open(str(path))
        text = ""
        for page in doc:
            text += page.get_text() or ""
        preview_path = None
        if len(doc) > 0:
            page = doc[0]
            pix = page.get_pixmap()
            previews_dir.mkdir(parents=True, exist_ok=True)
            preview_path = previews_dir / (path.stem + "_page1.png")
            pix.save(str(preview_path))
        doc.close()
        return text, preview_path
    except Exception as e:
        print(f"⚠️  PDF parse failed on {path.name}: {e}")
        return "", None

def _extract_docx(path: Path) -> str:
    if not DocxDocument:
        print("⚠️  python-docx not available for DOCX extract (should already be installed).")
        return ""
    try:
        doc = DocxDocument(str(path))
        parts = []
        for p in doc.paragraphs:
            if p.text:
                parts.append(p.text)
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    if cell.text:
                        parts.append(cell.text)
        return "\n".join(parts)
    except Exception as e:
        print(f"⚠️  DOCX extract failed on {path.name}: {e}")
        return ""

def _extract_txt(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8", errors="ignore")
    except Exception as e:
        print(f"⚠️  TXT read failed on {path.name}: {e}")
        return ""

def extract_text_and_preview(path: Path, previews_dir: Path) -> tuple[str, Optional[Path]]:
    ext = path.suffix.lower()
    if ext in {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"}:
        return _ocr_image(path), path  # preview is original image
    if ext == ".pdf":
        return _extract_pdf(path, previews_dir)
    if ext == ".docx":
        return _extract_docx(path), None
    if ext == ".txt":
        return _extract_txt(path), None
    print(f"⚠️  Unsupported evidence type: {ext}")
    return "", None


# ------------------------ Main --------------------------
def main():
    configure_tesseract() 

    user_id = input("Enter your username: ").strip()

    strategies = load_strategies()
    if not strategies:
        print("No strategies found.")
        return

    def desc(s):
        try:
            return s.description()
        except Exception:
            return ""

    names_for_menu = [f"{s.name} — {desc(s)}" for s in strategies]
    chosen_label = choose_one("\nAvailable strategies:", names_for_menu)
    chosen_name = chosen_label.split(" — ")[0]
    chosen_strategy = next((s for s in strategies if s.name == chosen_name), None)
    if not chosen_strategy:
        print("Not a valid selection. Exiting.")
        return

    print(f"\n📋 Strategy selected: {chosen_strategy.name}\n")

    evidence_dir = Path("evidence")
    evidence_dir.mkdir(parents=True, exist_ok=True)

    evidence_path = choose_one_evidence(evidence_dir)
    if not evidence_path:
        print("No evidence selected.")
        return

    out_dir = Path("reports_out")
    out_dir.mkdir(parents=True, exist_ok=True)
    previews_dir = Path("previews")  
    template_path = "templates/report_template.docx"

    print(f"\n📄  Evidence: {evidence_path.name}")
    text, preview = extract_text_and_preview(evidence_path, previews_dir)
    preview_text = (text[:300] + "…") if len(text) > 300 else text
    print("📝 Extracted Text:", preview_text)

    generated = []

    try:
        if hasattr(chosen_strategy, "emit_hits"):
            rows = chosen_strategy.emit_hits(text, source_file=evidence_path.name) or []
            if not rows:
                print("No findings for this strategy.")
            for idx, r in enumerate(rows, start=1):
                uid = _safe_uid(user_id)
                data = {
                    "UniqueID": uid,
                    "UserID": user_id,
                    "Evidence": str(evidence_path),
                    "Evidence Preview": str(preview) if preview else "",
                    "Strategy": chosen_strategy.name,
                    "TestID": r.get("test_id", ""),
                    "Sub-Strategy": r.get("sub_strategy", ""),
                    "ML Level": r.get("detected_level", ""),
                    "Pass/Fail": r.get("pass_fail", ""),
                    "Priority": r.get("priority", ""),
                    "Recommendation": r.get("recommendation", ""),
                    "Evidence Extract": "; ".join(r.get("evidence", [])),
                    "Description": r.get("description", ""),
                    "Confidence": r.get("confidence", ""),
                }
                pdf = generate_pdf(
                    data,
                    template_path=template_path,
                    output_dir=str(out_dir),
                    base_dir="."
                )
                print(f"   ✅ {chosen_strategy.name} → {pdf}")
                generated.append(pdf)
        else:
            hits = chosen_strategy.match(text) or []
            if not hits:
                print("No findings for this strategy.")
            else:
                uid = _safe_uid(f"{user_id}-{chosen_strategy.name}-{evidence_path.stem}-HITS")
                data = {
                    "UniqueID": uid,
                    "UserID": user_id,
                    "Evidence": str(evidence_path),
                    "Evidence Preview": str(preview) if preview else "",
                    "Strategy": chosen_strategy.name,
                    "TestID": "",
                    "Sub-Strategy": "",
                    "ML Level": "",
                    "Pass/Fail": "",
                    "Priority": "",
                    "Recommendation": "",
                    "Evidence Extract": ", ".join(hits),
                }
                pdf = generate_pdf(
                    data,
                    template_path=template_path,
                    output_dir=str(out_dir),
                    base_dir="."
                )
                print(f"   ✅ {chosen_strategy.name} → {pdf}")
                generated.append(pdf)
    except Exception as e:
        print(f"   ❌ {chosen_strategy.name} failed: {e}")

    if generated:
        print("\n📄 Generated reports:")
        for p in generated:
            print(" -", p)
    else:
        print("\nNo findings → no reports generated.")


# -------------------- utilities --------------------
def _safe_uid(s: str) -> str:
    keep = []
    for ch in s:
        if ch.isalnum() or ch in ("-", "_"):
            keep.append(ch)
        else:
            keep.append("-")
    return "".join(keep).strip("-") or "report"


if __name__ == "__main__":
    main()
